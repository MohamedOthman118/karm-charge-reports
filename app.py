"""
KARM CHARGE — EV Charger Operations & Diagnostic Report Generator
─────────────────────────────────────────────────────────────────
Workflow:
  1. Copy the AI prompt → paste with raw field notes into any AI.
  2. AI returns JSON.
  3. Paste JSON here → click "Load JSON" → form gets pre-filled.
  4. Review and edit ANY field manually before generating.
  5. Click "Generate Report" → download polished .docx.
"""

import io
import json
from datetime import date, datetime

import streamlit as st
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS — dropdown options and empty row templates
# ═══════════════════════════════════════════════════════════════════════════════

SEVERITY_OPTIONS  = ["Critical", "High", "Medium", "Low", "Resolved", "Unknown"]
PHYS_STATUS_OPTS  = ["Operational", "Working", "Good", "Out of service", "Degraded", "Under repair", "Other"]
DASH_STATUS_OPTS  = ["Online", "Offline", "Charging", "Available", "Faulted", "Unknown"]
PRIORITY_OPTIONS  = ["Critical", "High", "Medium", "Low"]
DIAG_STATUS_OPTS  = ["Open", "In Progress", "Resolved", "Escalated", "Closed"]

EMPTY_ASSET    = {"asset_id": "", "location": "", "physical_status": "Out of service",
                  "dashboard_status": "Offline", "severity": "Critical"}
EMPTY_DIAG     = {"issue": "", "affected_units": "", "symptom": "",
                  "root_cause": "", "action_taken": "", "status": "Open"}
EMPTY_ACTION   = {"issue": "", "next_step": "", "owner": "",
                  "priority": "Medium", "target_date": ""}
EMPTY_EVIDENCE = {"link": "", "asset_subject": "", "description": ""}

# ═══════════════════════════════════════════════════════════════════════════════
# AI PROMPT TEMPLATE (user copies this into ChatGPT / Claude / Gemini)
# ═══════════════════════════════════════════════════════════════════════════════

AI_PROMPT = """You are an EV Charging Infrastructure Operations Specialist working for KARM CHARGE.

I will paste raw field notes from a site visit. Your job is to convert them into a structured JSON object that follows the exact schema below. Do not invent data — if a field is not mentioned in the notes, leave it as an empty string, an empty list, or 0 as appropriate.

For "executive_summary", rephrase the raw notes into a professional, concise paragraph (3-6 sentences) focusing on technical diagnostics, root causes, and actions taken. Use clear business language — no bullets, no emojis, no markdown.

For "operational_objectives", extract the planned goals of the visit as short imperative statements (one per item).

Return ONLY valid JSON. No markdown code fences. No commentary before or after.

JSON SCHEMA:
{
  "site": "e.g. District 5 (Mall & Residential)",
  "visit_date": "YYYY-MM-DD",
  "report_id": "OPS-XXX-YYYY-MM-DD  (leave empty to auto-generate)",
  "prepared_by": "Mohamed Othman",
  "role": "Customer Ops Specialist",
  "personnel": [
    "Youssef Elasaeed — Karm Solar, Technical Team Member",
    "Sha7en field technicians — Warranty supplier"
  ],
  "executive_summary": "A professional 3-6 sentence summary in your own words.",
  "kpis": {
    "total_inspected": 19,
    "operational": 16,
    "offline": 3,
    "open_actions": 2
  },
  "operational_objectives": [
    "Diagnose relay-stuck fault on D5R-EVAC-01/03 with Sha7en warranty team",
    "Replace circuit breaker on D5M-EVAC-02/03",
    "Inspect all remaining chargers for operational status"
  ],
  "asset_status": [
    {
      "asset_id": "D5M-EVAC-01/1",
      "location": "D5M – Mall (LA7)",
      "physical_status": "Working",
      "dashboard_status": "Online",
      "severity": "Low"
    }
  ],
  "diagnostics": [
    {
      "issue": "Relay Stuck — D5R-EVAC-01/03",
      "affected_units": "D5R-EVAC-01/03",
      "symptom": "Charger offline; stuck relay fault reported.",
      "root_cause": "Awaiting Sha7en's report.",
      "action_taken": "Sha7en performed full diagnostic sequence; unit taken off-site for repair.",
      "status": "Open"
    }
  ],
  "open_action_items": [
    {
      "issue": "D5R-EVAC-01/03 offline — relay stuck; unit at Sha7en",
      "next_step": "Obtain written repair timeline from Sha7en",
      "owner": "Mohamed Othman",
      "priority": "Critical",
      "target_date": "2 business days"
    }
  ],
  "evidence_attachments": [
    {
      "link": "OPS-DS5-2026-05-19",
      "asset_subject": "D5R-EVAC-01/03, D5R-EVAC-01/04",
      "description": "Visit media: site conditions, breaker, charger handling"
    }
  ]
}

RAW FIELD NOTES:
<<<PASTE YOUR NOTES HERE>>>
"""

# ═══════════════════════════════════════════════════════════════════════════════
# SAMPLE JSON
# ═══════════════════════════════════════════════════════════════════════════════

SAMPLE_JSON = {
    "site": "District 5 (Mall & Residential)",
    "visit_date": "2026-05-19",
    "report_id": "OPS-DS5-2026-05-19",
    "prepared_by": "Mohamed Othman",
    "role": "Customer Ops Specialist",
    "personnel": [
        "Youssef Elasaeed — Karm Solar, Technical Team Member",
        "Sha7en field technicians — Warranty supplier (D5R diagnostic visit)",
    ],
    "executive_summary": (
        "On 19 May 2026, I visited District 5 Mall and Residential stations to address "
        "three reported issues across two sites. No issues were fully resolved on-site: "
        "the relay-stuck fault on D5R-EVAC-01/03 was diagnosed by Sha7en but requires "
        "off-site repair at their factory; the circuit-breaker replacement for "
        "D5M-EVAC-02/03 could not proceed due to a compatibility mismatch; and "
        "D5R-EVAC-01/04 remains online but under monitoring following a trip event "
        "the prior day. The immediate critical path is securing a repair timeline from "
        "Sha7en for D5R-EVAC-01/03 and raising a PO for a compatible circuit breaker."
    ),
    "kpis": {"total_inspected": 19, "operational": 16, "offline": 3, "open_actions": 2},
    "operational_objectives": [
        "Diagnose relay-stuck fault on D5R-EVAC-01/03 with Sha7en warranty team",
        "Replace circuit breaker on D5M-EVAC-02/03",
        "Assess D5R-EVAC-01/04 following prior-day trip issue",
        "Inspect all remaining D5M and D5R chargers for operational status",
    ],
    "asset_status": [
        {"asset_id": "D5M-EVAC-01/1", "location": "D5M – Mall (LA7)",
         "physical_status": "Working", "dashboard_status": "Online", "severity": "Low"},
        {"asset_id": "D5M-EVAC-02/03", "location": "D5M – Mall",
         "physical_status": "Out of service", "dashboard_status": "Offline", "severity": "Critical"},
        {"asset_id": "D5R-EVAC-01/03", "location": "D5R – Residential",
         "physical_status": "Out of service", "dashboard_status": "Offline", "severity": "Critical"},
        {"asset_id": "D5R-EVAC-01/04", "location": "D5R – Residential",
         "physical_status": "Working", "dashboard_status": "Online", "severity": "Medium"},
    ],
    "diagnostics": [
        {"issue": "Relay Stuck — D5R-EVAC-01/03",
         "affected_units": "D5R-EVAC-01/03",
         "symptom": "Charger offline; stuck relay fault reported.",
         "root_cause": "Awaiting Sha7en's report.",
         "action_taken": "Sha7en performed full diagnostic sequence (remote → tester → disassembly); unit taken off-site for repair.",
         "status": "Open"},
        {"issue": "Circuit Breaker Failure — D5M-EVAC-02/03",
         "affected_units": "D5M-EVAC-02/03",
         "symptom": "Charger fully offline; circuit breaker confirmed failed.",
         "root_cause": "Breaker hardware failure.",
         "action_taken": "Attempted to reuse a spare breaker from D5R — incompatible vertical mount.",
         "status": "Open"},
    ],
    "open_action_items": [
        {"issue": "D5R-EVAC-01/03 offline — relay stuck; unit at Sha7en",
         "next_step": "Obtain written repair timeline from Sha7en; escalate if needed",
         "owner": "Mohamed Othman", "priority": "Critical", "target_date": "2 business days"},
        {"issue": "D5M-EVAC-02/03 offline — incompatible breaker, unit still down",
         "next_step": "Raise PO for correct vertical-mount circuit breaker; schedule install",
         "owner": "Mohamed Othman + Noaman", "priority": "Critical", "target_date": "Once PO is approved"},
    ],
    "evidence_attachments": [
        {"link": "OPS-DS5-2026-05-19",
         "asset_subject": "D5R-EVAC-01/03, D5R-EVAC-01/04",
         "description": "Visit media: site conditions, breaker, charger handling"},
    ],
}

# ═══════════════════════════════════════════════════════════════════════════════
# JSON PARSING / NORMALISATION
# ═══════════════════════════════════════════════════════════════════════════════

def _strip_markdown_fences(raw: str) -> str:
    s = raw.strip()
    if s.startswith("```"):
        s = s.split("\n", 1)[1] if "\n" in s else s
        if s.rstrip().endswith("```"):
            s = s.rstrip()[:-3]
    return s.strip()


def _safe_choice(value, options, default):
    if not value:
        return default
    for opt in options:
        if str(value).strip().lower() == opt.lower():
            return opt
    return default


def parse_json(raw_json: str) -> tuple[dict | None, list[str]]:
    warnings: list[str] = []
    cleaned = _strip_markdown_fences(raw_json)
    if not cleaned:
        return None, ["JSON is empty — paste the AI's response into the box above."]

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as e:
        return None, [f"❌ Invalid JSON: {e.msg} (line {e.lineno}, col {e.colno})"]
    if not isinstance(data, dict):
        return None, ["❌ Top-level JSON must be an object {...}."]

    site = str(data.get("site", "")).strip()
    raw_date = str(data.get("visit_date", "")).strip()
    try:
        visit_dt = datetime.strptime(raw_date, "%Y-%m-%d").date()
    except (ValueError, TypeError):
        visit_dt = date.today()
        if raw_date:
            warnings.append(f"⚠️ Could not parse visit_date '{raw_date}' — defaulted to today.")

    report_id = str(data.get("report_id", "")).strip()
    if not report_id:
        prefix = "".join(c for c in site.upper() if c.isalpha())[:3] or "GEN"
        report_id = f"OPS-{prefix}-{visit_dt.isoformat()}"

    kpis_raw = data.get("kpis", {}) or {}
    def _kpi(k):
        try:    return int(kpis_raw.get(k, 0))
        except: return 0
    kpis = {"total_inspected": _kpi("total_inspected"), "operational": _kpi("operational"),
            "offline": _kpi("offline"), "open_actions": _kpi("open_actions")}

    personnel = [str(p).strip() for p in (data.get("personnel") or []) if str(p).strip()]
    if not personnel: personnel = [""]

    objectives = [str(o).strip() for o in (data.get("operational_objectives") or []) if str(o).strip()]
    if not objectives: objectives = [""]

    assets = []
    for a in (data.get("asset_status") or []):
        if not isinstance(a, dict): continue
        assets.append({
            "asset_id":         str(a.get("asset_id", "")).strip(),
            "location":         str(a.get("location", "")).strip(),
            "physical_status":  _safe_choice(a.get("physical_status"),  PHYS_STATUS_OPTS, "Out of service"),
            "dashboard_status": _safe_choice(a.get("dashboard_status"), DASH_STATUS_OPTS, "Offline"),
            "severity":         _safe_choice(a.get("severity"),         SEVERITY_OPTIONS, "Critical"),
        })
    if not assets: assets = [dict(EMPTY_ASSET)]

    diags = []
    for d in (data.get("diagnostics") or []):
        if not isinstance(d, dict): continue
        diags.append({
            "issue":          str(d.get("issue", "")).strip(),
            "affected_units": str(d.get("affected_units", "")).strip(),
            "symptom":        str(d.get("symptom", "")).strip(),
            "root_cause":     str(d.get("root_cause", "")).strip(),
            "action_taken":   str(d.get("action_taken", "")).strip(),
            "status":         _safe_choice(d.get("status"), DIAG_STATUS_OPTS, "Open"),
        })
    if not diags: diags = [dict(EMPTY_DIAG)]

    actions = []
    for a in (data.get("open_action_items") or []):
        if not isinstance(a, dict): continue
        actions.append({
            "issue":       str(a.get("issue", "")).strip(),
            "next_step":   str(a.get("next_step", "")).strip(),
            "owner":       str(a.get("owner", "")).strip(),
            "priority":    _safe_choice(a.get("priority"), PRIORITY_OPTIONS, "Medium"),
            "target_date": str(a.get("target_date", "")).strip(),
        })
    if not actions: actions = [dict(EMPTY_ACTION)]

    evidence = []
    for e in (data.get("evidence_attachments") or []):
        if not isinstance(e, dict): continue
        evidence.append({
            "link":          str(e.get("link", "")).strip(),
            "asset_subject": str(e.get("asset_subject", "")).strip(),
            "description":   str(e.get("description", "")).strip(),
        })
    if not evidence: evidence = [dict(EMPTY_EVIDENCE)]

    return {
        "site": site, "visit_date": visit_dt, "report_id": report_id,
        "prepared_by": str(data.get("prepared_by", "Mohamed Othman")).strip() or "Mohamed Othman",
        "role":        str(data.get("role", "Customer Ops Specialist")).strip() or "Customer Ops Specialist",
        "personnel":             personnel,
        "executive_summary":     str(data.get("executive_summary", "")).strip(),
        "kpis":                  kpis,
        "operational_objectives": objectives,
        "asset_status":          assets,
        "diagnostics":           diags,
        "open_action_items":     actions,
        "evidence_attachments":  evidence,
    }, warnings


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX COLOUR PALETTE & STYLING HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

NAVY_BLUE  = RGBColor(0x1B, 0x3A, 0x6B)    # primary brand colour — matches reference
BLACK_TEXT = RGBColor(0x00, 0x00, 0x00)
WHITE_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TEXT  = RGBColor(0x66, 0x66, 0x66)

NAVY_FILL    = "1B3A6B"   # for headers, title-page label column
LIGHT_FILL   = "E8EEF7"   # very light blue-grey for KV label column
ZEBRA_FILL   = "F5F7FB"   # almost-white zebra stripe
SUMMARY_FILL = "EAF1FA"   # AT A GLANCE box
RULE_GREY    = "BFBFBF"


def _set_cell_fill(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_run(p, text, size_pt=11, color=None, italic=False):
    r = p.add_run(text); r.bold = True; r.italic = italic; r.font.size = Pt(size_pt)
    if color: r.font.color.rgb = color
    return r


def _normal_run(p, text, size_pt=10, color=None, italic=False):
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size_pt)
    if color: r.font.color.rgb = color
    return r


def _add_section_heading(doc, text):
    """Navy section heading with a thin navy underline (matches reference)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(14); run.font.color.rgb = NAVY_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8"); bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), "1B3A6B")
    pBdr.append(bot); pPr.append(pBdr)


def _add_paragraph_border(paragraph, edge, color_hex=RULE_GREY, size="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bdr = OxmlElement(f"w:{edge}")
    bdr.set(qn("w:val"), "single"); bdr.set(qn("w:sz"), size); bdr.set(qn("w:space"), "1"); bdr.set(qn("w:color"), color_hex)
    pBdr.append(bdr); pPr.append(pBdr)


def _remove_cell_borders(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None: tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}"); b.set(qn("w:val"), "nil"); tcBorders.append(b)
    tcPr.append(tcBorders)


# ── Tables ──────────────────────────────────────────────────────────────────

def _add_title_page_table(doc, rows):
    """
    The big metadata table on the title page — navy left column with WHITE BOLD UPPERCASE labels,
    matches the reference document's first page exactly.
    """
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(3.5)

    for ri, (label, value) in enumerate(rows):
        lc, vc = table.rows[ri].cells
        _set_cell_fill(lc, NAVY_FILL)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4); lp.paragraph_format.space_after = Pt(4)
        _bold_run(lp, label.upper(), size_pt=10, color=WHITE_TEXT)

        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4); vp.paragraph_format.space_after = Pt(4)
        _normal_run(vp, str(value), size_pt=10)


def _add_kv_table(doc, rows):
    """Two-column label/value table for the Document Control section (light-grey style)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (label, value) in enumerate(rows):
        lc, vc = table.rows[ri].cells
        _set_cell_fill(lc, LIGHT_FILL)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(3); lp.paragraph_format.space_after = Pt(3)
        _bold_run(lp, label, size_pt=10, color=NAVY_BLUE)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3); vp.paragraph_format.space_after = Pt(3)
        _normal_run(vp, str(value), size_pt=10)


def _add_data_table(doc, columns, data_rows):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(columns))
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, col in enumerate(columns):
        cell = table.rows[0].cells[ci]; _set_cell_fill(cell, NAVY_FILL)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        _bold_run(p, col, size_pt=10, color=WHITE_TEXT)
    for ri, data_row in enumerate(data_rows):
        fill = ZEBRA_FILL if ri % 2 == 0 else "FFFFFF"
        for ci, value in enumerate(data_row):
            cell = table.rows[ri + 1].cells[ci]; _set_cell_fill(cell, fill)
            p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            _normal_run(p, str(value), size_pt=10)


def _add_kpi_table(doc, kpis):
    """KPI strip — 4 cards. Each = (label, value, hex_fill)."""
    table = doc.add_table(rows=1, cols=len(kpis))
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, (label, value, fill_hex) in enumerate(kpis):
        cell = table.rows[0].cells[ci]; _set_cell_fill(cell, fill_hex)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(10)
        vr = p.add_run(f"{value}\n"); vr.bold = True; vr.font.size = Pt(22); vr.font.color.rgb = WHITE_TEXT
        lr = p.add_run(label.upper()); lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = WHITE_TEXT


def _add_at_a_glance_table(doc, summary_text):
    """Soft light-blue summary card (matches reference)."""
    table = doc.add_table(rows=2, cols=1); table.style = "Table Grid"
    # Header bar (white text on navy)
    tc = table.rows[0].cells[0]; _set_cell_fill(tc, NAVY_FILL)
    tp = tc.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_before = Pt(4); tp.paragraph_format.space_after = Pt(4)
    _bold_run(tp, "AT A GLANCE", size_pt=10, color=WHITE_TEXT)
    # Body
    bc = table.rows[1].cells[0]; _set_cell_fill(bc, SUMMARY_FILL)
    bp = bc.paragraphs[0]
    bp.paragraph_format.space_before = Pt(6); bp.paragraph_format.space_after = Pt(6)
    _normal_run(bp, summary_text or "(No summary provided)", size_pt=10)


# ═══════════════════════════════════════════════════════════════════════════════
# HEADER & FOOTER (repeats on every page after title)
# ═══════════════════════════════════════════════════════════════════════════════

def _add_page_number_field(run):
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr     = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end   = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin); run._element.append(instr); run._element.append(fld_end)


def _apply_header_footer(doc, report_id):
    section = doc.sections[0]

    # ── HEADER (3-column borderless table) ──────────────────────────────────
    header = section.header
    header.paragraphs[0].text = ""
    h_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    h_table.autofit = False; h_table.allow_autofit = False
    h_table.columns[0].width = Inches(3.6)
    h_table.columns[1].width = Inches(1.5)
    h_table.columns[2].width = Inches(1.4)

    l_cell, m_cell, r_cell = h_table.rows[0].cells
    for c in (l_cell, m_cell, r_cell): _remove_cell_borders(c)

    l_cell.width = Inches(3.6)
    lp = l_cell.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rb = lp.add_run("KARM CHARGE"); rb.bold = True; rb.font.size = Pt(8); rb.font.color.rgb = NAVY_BLUE
    rs = lp.add_run("  —  ");        rs.font.size = Pt(8); rs.font.color.rgb = GREY_TEXT
    rt = lp.add_run("EV Charger Operations & Diagnostic Report")
    rt.bold = True; rt.font.size = Pt(8); rt.font.color.rgb = NAVY_BLUE

    m_cell.width = Inches(1.5)
    mp = m_cell.paragraphs[0]; mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = mp.add_run("Internal – Operations"); rm.font.size = Pt(8); rm.font.color.rgb = GREY_TEXT

    r_cell.width = Inches(1.4)
    rp = r_cell.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(report_id); rr.font.size = Pt(8); rr.font.color.rgb = GREY_TEXT

    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(0); rule_p.paragraph_format.space_after = Pt(0)
    _add_paragraph_border(rule_p, edge="bottom", color_hex="1B3A6B", size="6")

    # ── FOOTER (top rule + 2-column borderless table) ────────────────────────
    footer = section.footer
    rule_f = footer.paragraphs[0]; rule_f.text = ""
    rule_f.paragraph_format.space_before = Pt(0); rule_f.paragraph_format.space_after = Pt(0)
    _add_paragraph_border(rule_f, edge="top")

    f_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    f_table.autofit = False; f_table.allow_autofit = False
    f_table.columns[0].width = Inches(5.5); f_table.columns[1].width = Inches(1.0)
    fl, fr = f_table.rows[0].cells
    for c in (fl, fr): _remove_cell_borders(c)

    fl.width = Inches(5.5)
    flp = fl.paragraphs[0]; flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fl_run = flp.add_run("Karm Charge  ·  Operations Report  ·  Internal")
    fl_run.italic = True; fl_run.font.size = Pt(8); fl_run.font.color.rgb = GREY_TEXT

    fr.width = Inches(1.0)
    frp = fr.paragraphs[0]; frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pn = frp.add_run(); pn.font.size = Pt(8); pn.font.color.rgb = GREY_TEXT
    _add_page_number_field(pn)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_docx(data: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    _apply_header_footer(doc, data.get("report_id", ""))

    visit_str = (data["visit_date"].strftime("%A – %d %B %Y")
                 if hasattr(data["visit_date"], "strftime") else str(data["visit_date"]))

    # ═════════════════════════════════════════════════════════════════════════
    # TITLE PAGE — centred KARM CHARGE block + metadata table + page break
    # ═════════════════════════════════════════════════════════════════════════

    # spacer
    for _ in range(4):
        doc.add_paragraph()

    # KARM CHARGE — big navy title
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("KARM CHARGE")
    tr.bold = True; tr.font.size = Pt(32); tr.font.color.rgb = NAVY_BLUE

    doc.add_paragraph()

    # Sub-title (BLACK, bold) — matches reference design
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("EV Charger Operations & Diagnostic Report")
    sr.bold = True; sr.font.size = Pt(20); sr.font.color.rgb = BLACK_TEXT

    # Tagline
    gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gr = gp.add_run("Site visit operations report")
    gr.italic = True; gr.font.size = Pt(11); gr.font.color.rgb = GREY_TEXT

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table (centred, navy left column)
    _add_title_page_table(doc, [
        ("Site / District", data.get("site", "")),
        ("Date of Visit",   visit_str),
        ("Report ID",       data["report_id"]),
        ("Prepared By",     f"{data.get('prepared_by', 'Mohamed Othman')}, {data.get('role', '')}"),
        ("Version",         "v1.0"),
        ("Classification",  "Internal – Operations"),
    ])

    # Page break before content
    pb_p = doc.add_paragraph()
    pb_p.add_run().add_break(WD_BREAK.PAGE)

    # ═════════════════════════════════════════════════════════════════════════
    # PAGE 2 ONWARDS — content sections
    # ═════════════════════════════════════════════════════════════════════════

    # ── Document Control (light-grey style, repeats key info) ────────────────
    _add_section_heading(doc, "Document Control")
    _add_kv_table(doc, [
        ("Report ID",       data["report_id"]),
        ("Site / District", data.get("site", "")),
        ("Date of Visit",   visit_str),
        ("Prepared By",     data.get("prepared_by", "Mohamed Othman")),
        ("Role",            data.get("role", "Customer Ops Specialist")),
        ("Classification",  "Internal – Operations"),
    ])
    doc.add_paragraph()

    # ── Accompanying Personnel ───────────────────────────────────────────────
    pers_p = doc.add_paragraph()
    pers_p.paragraph_format.space_before = Pt(4)
    _bold_run(pers_p, "Accompanying Personnel", size_pt=11, color=NAVY_BLUE)

    people = [p.strip() for p in data.get("personnel", []) if str(p).strip()]
    if people:
        for person in people:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_before = Pt(1); bp.paragraph_format.space_after = Pt(1)
            _normal_run(bp, person, size_pt=10)
    else:
        doc.add_paragraph("No accompanying personnel listed.").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    _add_section_heading(doc, "1. Executive Summary")
    _add_at_a_glance_table(doc, data.get("executive_summary", ""))
    doc.add_paragraph()

    # ── 2. Visit KPIs ────────────────────────────────────────────────────────
    _add_section_heading(doc, "2. Visit KPIs at a Glance")
    kpis = data.get("kpis", {})
    _add_kpi_table(doc, [
        ("Total Inspected",   kpis.get("total_inspected", 0), "1B3A6B"),   # navy
        ("Operational",       kpis.get("operational", 0),     "2E7D32"),   # forest green
        ("Offline / Faulted", kpis.get("offline", 0),         "C62828"),   # deep red
        ("Open Actions",      kpis.get("open_actions", 0),    "8C6D1F"),   # olive (matches reference)
    ])
    doc.add_paragraph()

    # ── 3. Operational Objectives ────────────────────────────────────────────
    _add_section_heading(doc, "3. Operational Objectives")
    objectives = [o.strip() for o in data.get("operational_objectives", []) if str(o).strip()]
    if objectives:
        for obj in objectives:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_before = Pt(1); bp.paragraph_format.space_after = Pt(1)
            _normal_run(bp, obj, size_pt=10)
    else:
        doc.add_paragraph("No operational objectives recorded.").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 4. Asset Status ──────────────────────────────────────────────────────
    _add_section_heading(doc, "4. Asset Status & Observations")
    asset_rows = [
        [a.get("asset_id", ""), a.get("location", ""),
         a.get("physical_status", ""), a.get("dashboard_status", ""),
         a.get("severity", "")]
        for a in data.get("asset_status", [])
        if any(str(v).strip() for v in a.values())
    ]
    if asset_rows:
        _add_data_table(doc,
            ["Asset ID", "Location", "Physical Status", "Dashboard Status", "Severity"],
            asset_rows)
    else:
        doc.add_paragraph("No asset records provided.").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 5. Diagnostics ───────────────────────────────────────────────────────
    _add_section_heading(doc, "5. Diagnostics & Actions Taken")
    diagnostics = [d for d in data.get("diagnostics", []) if any(str(v).strip() for v in d.values())]
    if not diagnostics:
        doc.add_paragraph("No diagnostic entries provided.").runs[0].font.size = Pt(10)
    else:
        for idx, d in enumerate(diagnostics, start=1):
            ih = doc.add_paragraph()
            ih.paragraph_format.space_before = Pt(8); ih.paragraph_format.space_after = Pt(2)
            _bold_run(ih, f"Issue {idx}: {d.get('issue', '—')}", size_pt=11, color=NAVY_BLUE)
            for label, key in [
                ("Affected Units", "affected_units"),
                ("Symptom",        "symptom"),
                ("Root Cause",     "root_cause"),
                ("Action Taken",   "action_taken"),
                ("Status",         "status"),
            ]:
                value = str(d.get(key, "")).strip()
                if value:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(1); bp.paragraph_format.space_after = Pt(1)
                    _bold_run(bp, f"{label}: ", size_pt=10)
                    _normal_run(bp, value, size_pt=10)
    doc.add_paragraph()

    # ── 6. Open Action Items ─────────────────────────────────────────────────
    _add_section_heading(doc, "6. Open Action Items")
    action_rows = []
    for a in data.get("open_action_items", []):
        if any(str(v).strip() for v in a.values()):
            action_rows.append([str(len(action_rows) + 1),
                a.get("issue", ""), a.get("next_step", ""),
                a.get("owner", ""), a.get("priority", ""), a.get("target_date", "")])
    if action_rows:
        _add_data_table(doc,
            ["#", "Issue", "Next Step", "Owner", "Priority", "Target Date"],
            action_rows)
    else:
        doc.add_paragraph("No open action items logged.").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 9. Evidence & Attachments ────────────────────────────────────────────
    _add_section_heading(doc, "9. Evidence & Attachments")
    ev_rows = []
    for e in data.get("evidence_attachments", []):
        if any(str(v).strip() for v in e.values()):
            ev_rows.append([str(len(ev_rows) + 1),
                e.get("link", ""), e.get("asset_subject", ""), e.get("description", "")])
    if ev_rows:
        _add_data_table(doc,
            ["#", "File / Link", "Asset / Subject", "Description"],
            ev_rows)
    else:
        doc.add_paragraph("No evidence or attachments listed.").runs[0].font.size = Pt(10)

    # ── Save ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════════

def init_state():
    defaults = {
        "form_version":            0,
        "site":                    "",
        "visit_date":              date.today(),
        "report_id":               "",
        "prepared_by":             "Mohamed Othman",
        "role":                    "Customer Ops Specialist",
        "personnel":               [""],
        "executive_summary":       "",
        "kpi_total":               0,
        "kpi_operational":         0,
        "kpi_offline":             0,
        "kpi_open_actions":        0,
        "operational_objectives":  [""],
        "asset_rows":              [dict(EMPTY_ASSET)],
        "diag_rows":               [dict(EMPTY_DIAG)],
        "action_rows":             [dict(EMPTY_ACTION)],
        "evidence_rows":           [dict(EMPTY_EVIDENCE)],
        "json_text":               "",
        "docx_bytes":              None,
        "docx_filename":           None,
        "load_warnings":           [],
        "show_json_panel":         True,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def load_from_json(raw_json: str) -> bool:
    parsed, warnings = parse_json(raw_json)
    if parsed is None:
        st.session_state["load_warnings"] = warnings
        return False

    st.session_state["form_version"]            += 1
    st.session_state["site"]                     = parsed["site"]
    st.session_state["visit_date"]               = parsed["visit_date"]
    st.session_state["report_id"]                = parsed["report_id"]
    st.session_state["prepared_by"]              = parsed["prepared_by"]
    st.session_state["role"]                     = parsed["role"]
    st.session_state["personnel"]                = parsed["personnel"]
    st.session_state["executive_summary"]        = parsed["executive_summary"]
    st.session_state["kpi_total"]                = parsed["kpis"]["total_inspected"]
    st.session_state["kpi_operational"]          = parsed["kpis"]["operational"]
    st.session_state["kpi_offline"]              = parsed["kpis"]["offline"]
    st.session_state["kpi_open_actions"]         = parsed["kpis"]["open_actions"]
    st.session_state["operational_objectives"]   = parsed["operational_objectives"]
    st.session_state["asset_rows"]               = parsed["asset_status"]
    st.session_state["diag_rows"]                = parsed["diagnostics"]
    st.session_state["action_rows"]              = parsed["open_action_items"]
    st.session_state["evidence_rows"]            = parsed["evidence_attachments"]
    st.session_state["load_warnings"]            = warnings
    st.session_state["docx_bytes"]               = None
    return True


def collect_form_data() -> dict:
    return {
        "site":                    st.session_state["site"],
        "visit_date":              st.session_state["visit_date"],
        "report_id":               st.session_state["report_id"] or
                                   f"OPS-GEN-{st.session_state['visit_date'].isoformat()}",
        "prepared_by":             st.session_state["prepared_by"],
        "role":                    st.session_state["role"],
        "personnel":               st.session_state["personnel"],
        "executive_summary":       st.session_state["executive_summary"],
        "kpis": {
            "total_inspected": st.session_state["kpi_total"],
            "operational":     st.session_state["kpi_operational"],
            "offline":         st.session_state["kpi_offline"],
            "open_actions":    st.session_state["kpi_open_actions"],
        },
        "operational_objectives":  st.session_state["operational_objectives"],
        "asset_status":            st.session_state["asset_rows"],
        "diagnostics":             st.session_state["diag_rows"],
        "open_action_items":       st.session_state["action_rows"],
        "evidence_attachments":    st.session_state["evidence_rows"],
    }


# ═══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="KARM CHARGE — Report Generator",
        page_icon="⚡",
        layout="wide",
        initial_sidebar_state="collapsed",
    )
    init_state()
    v = st.session_state["form_version"]

    # ── Banner (navy to match report) ────────────────────────────────────────
    st.markdown(
        """
        <div style='background:linear-gradient(90deg,#1B3A6B,#2E5C9C);
                    padding:18px 24px;border-radius:8px;margin-bottom:12px'>
          <h2 style='color:white;margin:0'>⚡ KARM CHARGE</h2>
          <p style='color:#d0eaff;margin:4px 0 0'>
            EV Charger Operations &amp; Diagnostic Report Generator
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        **How to use this tool**
        1. Copy the **AI prompt** below.
        2. Paste it into any AI (ChatGPT, Claude, Gemini) along with your raw field notes.
        3. AI returns **JSON** → paste it here → **Load JSON**.
        4. Review and edit any field, then **Generate Report**.
        """
    )

    st.divider()

    # ── Step 1 — AI prompt ───────────────────────────────────────────────────
    with st.expander("🤖 **Step 1** — Need the AI prompt? Click to copy", expanded=False):
        st.caption("Click the copy icon in the top-right of the box.")
        st.code(AI_PROMPT, language="markdown")

    # ── Step 2 — Paste JSON ──────────────────────────────────────────────────
    with st.expander("📥 **Step 2** — Paste JSON to auto-fill the form below",
                     expanded=st.session_state["show_json_panel"]):
        st.caption("Paste your AI's JSON, click **Load JSON**, every field below gets pre-filled. Then edit freely.")

        col_a, col_b = st.columns([1, 1])
        with col_a:
            if st.button("📝 Load sample JSON", use_container_width=True):
                st.session_state["json_text"] = json.dumps(SAMPLE_JSON, indent=2, ensure_ascii=False)
                st.rerun()
        with col_b:
            if st.button("🗑️ Clear JSON box", use_container_width=True):
                st.session_state["json_text"] = ""
                st.session_state["load_warnings"] = []
                st.rerun()

        json_text = st.text_area(
            "JSON",
            value=st.session_state["json_text"],
            height=280,
            placeholder='{\n  "site": "...",\n  "visit_date": "2026-05-19",\n  ...\n}',
            label_visibility="collapsed",
            key=f"json_text_widget_v{v}",
        )
        st.session_state["json_text"] = json_text

        if st.button("⬇️ **Load JSON into form below**", type="primary", use_container_width=True):
            if load_from_json(st.session_state["json_text"]):
                st.session_state["show_json_panel"] = False
                st.success("✅ Form populated from JSON. Scroll down to review and edit.")
                st.rerun()
            else:
                for w in st.session_state["load_warnings"]:
                    st.error(w)

        for w in st.session_state["load_warnings"]:
            if w.startswith("⚠️") or w.startswith("ℹ️"):
                st.warning(w)

    st.divider()
    st.markdown("### ✏️ **Step 3** — Review & edit before generating")

    # ── 📍 Visit Details ─────────────────────────────────────────────────────
    st.subheader("📍 Visit Details")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state["site"] = st.text_input(
            "Site / District", value=st.session_state["site"], key=f"site_v{v}",
            placeholder="e.g. District 5 (Mall & Residential)")
    with c2:
        st.session_state["visit_date"] = st.date_input(
            "Date of Visit", value=st.session_state["visit_date"], key=f"visit_date_v{v}")

    c3, c4, c5 = st.columns(3)
    with c3:
        st.session_state["report_id"] = st.text_input(
            "Report ID (leave blank to auto-generate)",
            value=st.session_state["report_id"], key=f"report_id_v{v}")
    with c4:
        st.session_state["prepared_by"] = st.text_input(
            "Prepared By", value=st.session_state["prepared_by"], key=f"prep_v{v}")
    with c5:
        st.session_state["role"] = st.text_input(
            "Role", value=st.session_state["role"], key=f"role_v{v}")

    st.divider()

    # ── 👥 Personnel ─────────────────────────────────────────────────────────
    st.subheader("👥 Accompanying Personnel")
    personnel = st.session_state["personnel"]
    for i in range(len(personnel)):
        col_p, col_x = st.columns([10, 1])
        with col_p:
            personnel[i] = st.text_input(
                f"Person {i+1}", value=personnel[i], key=f"pers_{i}_v{v}",
                label_visibility="collapsed",
                placeholder="e.g. Eng. Ahmed Nour — Solargy, Engineer")
        with col_x:
            if st.button("✖", key=f"pers_del_{i}_v{v}"):
                if len(personnel) > 1: personnel.pop(i)
                else: personnel[0] = ""
                st.rerun()
    if st.button("➕ Add Person", key=f"add_pers_v{v}"):
        personnel.append(""); st.rerun()

    st.divider()

    # ── 📝 Executive Summary ─────────────────────────────────────────────────
    st.subheader("📝 Executive Summary")
    st.caption("Goes into the AT A GLANCE box.")
    st.session_state["executive_summary"] = st.text_area(
        "Executive Summary", value=st.session_state["executive_summary"],
        height=180, key=f"summary_v{v}", label_visibility="collapsed",
        placeholder="Write or paste a 3-6 sentence professional summary of the visit.")

    st.divider()

    # ── 📊 KPIs ──────────────────────────────────────────────────────────────
    st.subheader("📊 Visit KPIs")
    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.session_state["kpi_total"] = st.number_input(
            "Total Inspected", min_value=0, step=1,
            value=st.session_state["kpi_total"], key=f"kpi1_v{v}")
    with k2:
        st.session_state["kpi_operational"] = st.number_input(
            "Operational", min_value=0, step=1,
            value=st.session_state["kpi_operational"], key=f"kpi2_v{v}")
    with k3:
        st.session_state["kpi_offline"] = st.number_input(
            "Offline / Faulted", min_value=0, step=1,
            value=st.session_state["kpi_offline"], key=f"kpi3_v{v}")
    with k4:
        st.session_state["kpi_open_actions"] = st.number_input(
            "Open Actions", min_value=0, step=1,
            value=st.session_state["kpi_open_actions"], key=f"kpi4_v{v}")

    st.divider()

    # ── 🎯 Operational Objectives ────────────────────────────────────────────
    st.subheader("🎯 Operational Objectives")
    st.caption("Bullet list of the planned goals for this visit.")
    objectives = st.session_state["operational_objectives"]
    for i in range(len(objectives)):
        col_o, col_x = st.columns([10, 1])
        with col_o:
            objectives[i] = st.text_input(
                f"Objective {i+1}", value=objectives[i], key=f"obj_{i}_v{v}",
                label_visibility="collapsed",
                placeholder="e.g. Diagnose relay-stuck fault on D5R-EVAC-01/03")
        with col_x:
            if st.button("✖", key=f"obj_del_{i}_v{v}"):
                if len(objectives) > 1: objectives.pop(i)
                else: objectives[0] = ""
                st.rerun()
    if st.button("➕ Add Objective", key=f"add_obj_v{v}"):
        objectives.append(""); st.rerun()

    st.divider()

    # ── 🔌 Asset Status ──────────────────────────────────────────────────────
    st.subheader("🔌 Asset Status & Observations")
    asset_rows = st.session_state["asset_rows"]
    for i, row in enumerate(asset_rows):
        with st.container(border=True):
            tl, tr = st.columns([10, 1])
            with tl: st.markdown(f"**Asset row {i+1}**")
            with tr:
                if st.button("✖", key=f"asset_del_{i}_v{v}"):
                    if len(asset_rows) > 1: asset_rows.pop(i)
                    else: asset_rows[0] = dict(EMPTY_ASSET)
                    st.rerun()
            c1, c2 = st.columns(2)
            row["asset_id"] = c1.text_input("Asset ID", value=row.get("asset_id", ""),
                key=f"a_id_{i}_v{v}", placeholder="D5M-EVAC-01/1")
            row["location"] = c2.text_input("Location", value=row.get("location", ""),
                key=f"a_loc_{i}_v{v}", placeholder="D5M – Mall (LA7)")
            c3, c4, c5 = st.columns(3)
            row["physical_status"] = c3.selectbox("Physical Status", options=PHYS_STATUS_OPTS,
                index=PHYS_STATUS_OPTS.index(row["physical_status"])
                      if row.get("physical_status") in PHYS_STATUS_OPTS else 3, key=f"a_ps_{i}_v{v}")
            row["dashboard_status"] = c4.selectbox("Dashboard Status", options=DASH_STATUS_OPTS,
                index=DASH_STATUS_OPTS.index(row["dashboard_status"])
                      if row.get("dashboard_status") in DASH_STATUS_OPTS else 1, key=f"a_ds_{i}_v{v}")
            row["severity"] = c5.selectbox("Severity", options=SEVERITY_OPTIONS,
                index=SEVERITY_OPTIONS.index(row["severity"])
                      if row.get("severity") in SEVERITY_OPTIONS else 0, key=f"a_sev_{i}_v{v}")
    if st.button("➕ Add Asset Row", key=f"add_asset_v{v}"):
        asset_rows.append(dict(EMPTY_ASSET)); st.rerun()

    st.divider()

    # ── 🔧 Diagnostics ───────────────────────────────────────────────────────
    st.subheader("🔧 Diagnostics & Actions Taken")
    diag_rows = st.session_state["diag_rows"]
    for i, row in enumerate(diag_rows):
        title = row.get("issue") or f"(untitled — issue {i+1})"
        with st.expander(f"🔧 Issue {i+1}: {title}", expanded=(i == 0)):
            row["issue"] = st.text_input("Issue Title", value=row.get("issue", ""),
                key=f"d_issue_{i}_v{v}", placeholder="e.g. Relay Stuck — D5R-EVAC-01/03")
            row["affected_units"] = st.text_input("Affected Units", value=row.get("affected_units", ""),
                key=f"d_units_{i}_v{v}")
            row["symptom"] = st.text_area("Symptom", value=row.get("symptom", ""),
                key=f"d_sym_{i}_v{v}", height=70)
            row["root_cause"] = st.text_area("Root Cause", value=row.get("root_cause", ""),
                key=f"d_rc_{i}_v{v}", height=70)
            row["action_taken"] = st.text_area("Action Taken", value=row.get("action_taken", ""),
                key=f"d_at_{i}_v{v}", height=70)
            row["status"] = st.selectbox("Status", options=DIAG_STATUS_OPTS,
                index=DIAG_STATUS_OPTS.index(row["status"])
                      if row.get("status") in DIAG_STATUS_OPTS else 0, key=f"d_stat_{i}_v{v}")
            if st.button("🗑️ Remove this issue", key=f"d_del_{i}_v{v}"):
                if len(diag_rows) > 1: diag_rows.pop(i)
                else: diag_rows[0] = dict(EMPTY_DIAG)
                st.rerun()
    if st.button("➕ Add Diagnostic", key=f"add_diag_v{v}"):
        diag_rows.append(dict(EMPTY_DIAG)); st.rerun()

    st.divider()

    # ── 📌 Open Action Items ─────────────────────────────────────────────────
    st.subheader("📌 Open Action Items")
    action_rows = st.session_state["action_rows"]
    for i, row in enumerate(action_rows):
        with st.container(border=True):
            tl, tr = st.columns([10, 1])
            with tl: st.markdown(f"**Action item {i+1}**")
            with tr:
                if st.button("✖", key=f"act_del_{i}_v{v}"):
                    if len(action_rows) > 1: action_rows.pop(i)
                    else: action_rows[0] = dict(EMPTY_ACTION)
                    st.rerun()
            c1, c2 = st.columns(2)
            row["issue"] = c1.text_input("Issue", value=row.get("issue", ""), key=f"ac_issue_{i}_v{v}")
            row["owner"] = c2.text_input("Owner", value=row.get("owner", ""), key=f"ac_own_{i}_v{v}")
            row["next_step"] = st.text_area("Next Step", value=row.get("next_step", ""),
                key=f"ac_ns_{i}_v{v}", height=60)
            c3, c4 = st.columns(2)
            row["priority"] = c3.selectbox("Priority", options=PRIORITY_OPTIONS,
                index=PRIORITY_OPTIONS.index(row["priority"])
                      if row.get("priority") in PRIORITY_OPTIONS else 2, key=f"ac_pri_{i}_v{v}")
            row["target_date"] = c4.text_input("Target Date", value=row.get("target_date", ""),
                key=f"ac_td_{i}_v{v}")
    if st.button("➕ Add Action Item", key=f"add_act_v{v}"):
        action_rows.append(dict(EMPTY_ACTION)); st.rerun()

    st.divider()

    # ── 📎 Evidence & Attachments ────────────────────────────────────────────
    st.subheader("📎 Evidence & Attachments")
    st.caption("Photos, dashboards, supporting files. Goes into Section 9 of the report.")
    evidence_rows = st.session_state["evidence_rows"]
    for i, row in enumerate(evidence_rows):
        with st.container(border=True):
            tl, tr = st.columns([10, 1])
            with tl: st.markdown(f"**Attachment {i+1}**")
            with tr:
                if st.button("✖", key=f"ev_del_{i}_v{v}"):
                    if len(evidence_rows) > 1: evidence_rows.pop(i)
                    else: evidence_rows[0] = dict(EMPTY_EVIDENCE)
                    st.rerun()
            c1, c2 = st.columns([2, 3])
            row["link"] = c1.text_input("File / Link", value=row.get("link", ""),
                key=f"ev_link_{i}_v{v}", placeholder="e.g. OPS-DS5-2026-05-19")
            row["asset_subject"] = c2.text_input("Asset / Subject", value=row.get("asset_subject", ""),
                key=f"ev_subj_{i}_v{v}", placeholder="e.g. D5R-EVAC-01/03")
            row["description"] = st.text_area("Description", value=row.get("description", ""),
                key=f"ev_desc_{i}_v{v}", height=60,
                placeholder="e.g. Visit media: site conditions, breaker, charger handling")
    if st.button("➕ Add Attachment", key=f"add_ev_v{v}"):
        evidence_rows.append(dict(EMPTY_EVIDENCE)); st.rerun()

    st.divider()

    # ── Step 4 — Generate & Download ─────────────────────────────────────────
    st.markdown("### 🚀 **Step 4** — Generate the report")

    col_gen, col_reset = st.columns([3, 1])
    with col_gen:
        if st.button("🚀 Generate Report", type="primary", use_container_width=True):
            try:
                data = collect_form_data()
                docx_bytes = build_docx(data)
                st.session_state["docx_bytes"]    = docx_bytes
                st.session_state["docx_filename"] = f"{data['report_id']}.docx"
                st.success("✅ Report generated! Download button below.")
            except Exception as exc:  # noqa: BLE001
                st.exception(exc)
    with col_reset:
        if st.button("🔄 Reset all fields", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()

    if st.session_state.get("docx_bytes"):
        st.download_button(
            label    = f"⬇️ Download  {st.session_state['docx_filename']}",
            data     = st.session_state["docx_bytes"],
            file_name= st.session_state["docx_filename"],
            mime     = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, type="primary")


if __name__ == "__main__":
    main()
